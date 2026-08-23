import pytest
import io
import json
import zipfile
import hashlib
import secrets
from fastapi.testclient import TestClient
import docx

from app.main import app
from app.db import Base, engine, SessionLocal
from app import models
from app.routers.auth import hash_password
from app.services.billing.bootstrap import ensure_plans_and_pricing_seeded

client = TestClient(app)


@pytest.fixture(autouse=True)
def setup_db():
    Base.metadata.create_all(bind=engine)
    yield


def create_test_tenant(db, username: str, org_id: str, role: str = "RESEARCHER"):
    user_email = f"{username}_{secrets.token_hex(4)}@test-univ.edu"
    user = db.query(models.User).filter(models.User.username == username).first()
    if not user:
        user = models.User(
            id=f"usr-{username}",
            username=username,
            email=user_email,
            hashed_password=hash_password("Password123!"),
            role="RESEARCHER",
            created_at="2026-08-22T00:00:00Z"
        )
        db.add(user)

    org = db.query(models.Organization).filter(models.Organization.id == org_id).first()
    if not org:
        org = models.Organization(
            id=org_id,
            name=f"University {org_id}",
            slug=f"slug-{org_id}",
            organization_type="UNIVERSITY",
            status="ACTIVE",
            owner_user_id=user.id,
            default_language="ar",
            data_region="sa",
            created_at="2026-08-22T00:00:00Z"
        )
        db.add(org)

    membership = db.query(models.OrganizationMembership).filter(
        models.OrganizationMembership.organization_id == org_id,
        models.OrganizationMembership.user_id == user.id
    ).first()
    if not membership:
        membership = models.OrganizationMembership(
            id=f"mbr-{username}-{secrets.token_hex(4)}",
            organization_id=org.id,
            user_id=user.id,
            role=role,
            status="ACTIVE",
            created_at="2026-08-22T00:00:00Z"
        )
        db.add(membership)

    ensure_plans_and_pricing_seeded(db)

    sub = db.query(models.Subscription).filter(models.Subscription.organization_id == org_id).first()
    if not sub:
        sub = models.Subscription(
            id=f"sub-{org_id}",
            organization_id=org.id,
            plan_id="pln-enterprise",
            status="ACTIVE",
            provider="MOCK",
            current_period_start="2026-08-22T00:00:00Z",
            current_period_end="2036-08-22T00:00:00Z",
            created_at="2026-08-22T00:00:00Z"
        )
        db.add(sub)
    else:
        sub.plan_id = "pln-enterprise"
        sub.status = "ACTIVE"

    db.commit()
    return user, org


def get_auth_headers(username: str, org_id: str):
    login_res = client.post("/api/auth/login", json={"username": username, "password": "Password123!"})
    token = login_res.json()["token"]
    return {
        "Authorization": f"Bearer {token}",
        "X-Organization-ID": org_id
    }


def seed_report_test_entities(db, user_a_id: str, org_a_id: str, editor_id: str):
    now_str = "2026-08-22T00:00:00Z"

    # Clean existing test entities if any to ensure fresh golden data
    db.query(models.ReviewComment).filter(models.ReviewComment.case_id == "case-rep-501").delete()
    db.query(models.ReviewSubmission).filter(models.ReviewSubmission.case_id == "case-rep-501").delete()
    db.query(models.ReviewerAssignment).filter(models.ReviewerAssignment.case_id == "case-rep-501").delete()
    db.query(models.PeerReviewRound).filter(models.PeerReviewRound.case_id == "case-rep-501").delete()
    db.query(models.PeerReviewCase).filter(models.PeerReviewCase.id == "case-rep-501").delete()

    db.query(models.PromotionAssetSelection).filter(models.PromotionAssetSelection.promotion_application_id == "app-rep-401").delete()
    db.query(models.PromotionApplication).filter(models.PromotionApplication.id == "app-rep-401").delete()
    db.query(models.PromotionPolicy).filter(models.PromotionPolicy.id == "pol-rep-1").delete()

    db.query(models.PrismaFlow).filter(models.PrismaFlow.id == "prisma-rep-301").delete()
    db.query(models.LiteratureStudy).filter(models.LiteratureStudy.projectId == "proj-rep-101").delete()
    db.query(models.Hypothesis).filter(models.Hypothesis.projectId == "proj-rep-101").delete()
    db.query(models.ResearchVariable).filter(models.ResearchVariable.projectId == "proj-rep-101").delete()
    db.query(models.ResearchProject).filter(models.ResearchProject.id == "proj-rep-101").delete()

    db.query(models.ScholarlyAsset).filter(models.ScholarlyAsset.id == "asset-rep-1").delete()
    db.query(models.UnifiedAcademicProfile).filter(models.UnifiedAcademicProfile.id == "prof-rep-601").delete()
    db.commit()

    # 1. Research Project
    proj_a = models.ResearchProject(
        id="proj-rep-101",
        organizationId=org_a_id,
        userId=user_a_id,
        titleAr="أثر تقنيات الذكاء الاصطناعي في جودة التعليم العالي",
        titleEn="Impact of AI Technologies on Higher Education Quality",
        studyDesign="QUANTITATIVE_EXPERIMENTAL",
        institutionAr="جامعة الملك سعود",
        institutionEn="King Saud University",
        departmentAr="قسم المناهج والتعليم",
        problemStatementAr="فجوة في قياس الأثر الفعلي لأدوات الذكاء الاصطناعي على التحصيل الأكاديمي.",
        descriptionAr="دراسة تجريبية مقترحة باستخدام مجموعات ضابطة وتجريبية.",
        sampleSettings={"marginOfError": 0.05, "expectedPower": 0.85, "expectedEffectSize": 0.65}
    )
    db.add(proj_a)

    var1 = models.ResearchVariable(id="var-rep-1", projectId=proj_a.id, nameAr="التحصيل المعرفي", nameEn="Cognitive Achievement", type="DEPENDENT", scale="INTERVAL")
    var2 = models.ResearchVariable(id="var-rep-2", projectId=proj_a.id, nameAr="استخدام بصيرة الذكي", nameEn="AI Tool Usage", type="INDEPENDENT", scale="NOMINAL")
    db.add_all([var1, var2])

    hypo1 = models.Hypothesis(id="hyp-rep-1", projectId=proj_a.id, textAr="توجد فروق ذات دلالة إحصائية لصالح المجموعة التجريبية", textEn="Statistically significant differences favoring experimental group", type="DIRECTIONAL")
    db.add(hypo1)

    # 2. Literature Studies
    st1 = models.LiteratureStudy(
        id="st-rep-1",
        projectId=proj_a.id,
        organizationId=org_a_id,
        author="Al-Salem et al., 2024",
        year=2024,
        sampleSize=120,
        effectSize=0.72,
        ciLower=0.45,
        ciUpper=0.99,
        source="manual",
        doi="10.1000/182",
        createdAt=now_str,
        updatedAt=now_str
    )
    st2 = models.LiteratureStudy(
        id="st-rep-2",
        projectId=proj_a.id,
        organizationId=org_a_id,
        author="Nasser, 2025",
        year=2025,
        sampleSize=95,
        effectSize=0.55,
        ciLower=0.25,
        ciUpper=0.85,
        source="manual",
        doi="10.1000/183",
        createdAt=now_str,
        updatedAt=now_str
    )
    db.add_all([st1, st2])

    # 3. PRISMA Flow
    prisma = models.PrismaFlow(
        id="prisma-rep-301",
        projectId=proj_a.id,
        organizationId=org_a_id,
        identified=450,
        duplicates=85,
        excludedScreening=280,
        excludedEligibility=40,
        source="manual",
        createdAt=now_str,
        updatedAt=now_str
    )
    db.add(prisma)

    # 4. Promotion Application
    pol = models.PromotionPolicy(
        id="pol-rep-1",
        organization_id=org_a_id,
        name_ar="لائحة ترقيات كبار الأساتذة 2026",
        name_en="Senior Professorship Promotion Policy 2026",
        version=1,
        target_rank="PROFESSOR",
        status="ACTIVE",
        created_at=now_str,
        updated_at=now_str
    )
    db.add(pol)

    app_prom = models.PromotionApplication(
        id="app-rep-401",
        organization_id=org_a_id,
        user_id=user_a_id,
        policy_id=pol.id,
        policy_version=1,
        target_rank="PROFESSOR",
        current_rank="ASSOCIATE_PROFESSOR",
        status="UNDER_REVIEW",
        total_calculated_points=79.0,
        readiness_percentage=98,
        created_at=now_str,
        updated_at=now_str
    )
    db.add(app_prom)

    # 5. Peer Review Case
    case = models.PeerReviewCase(
        id="case-rep-501",
        organization_id=org_a_id,
        owner_user_id=user_a_id,
        title_ar="منهجية القياس المتعدد في الذكاء الاصطناعي الطبي",
        title_en="Multi-Measurement Protocol in Medical AI",
        case_type="MANUSCRIPT",
        blind_type="DOUBLE_BLIND",
        status="IN_REVIEW",
        current_round_number=1,
        created_at=now_str,
        updated_at=now_str
    )
    db.add(case)

    rnd = models.PeerReviewRound(
        id="rnd-rep-1",
        case_id=case.id,
        round_number=1,
        manuscript_version=1,
        status="ACTIVE",
        created_at=now_str
    )
    db.add(rnd)

    asg = models.ReviewerAssignment(
        id="asg-rep-1",
        case_id=case.id,
        round_id=rnd.id,
        reviewer_user_id=editor_id,
        external_name="SECRET_REVIEWER_ALICE",
        status="SUBMITTED",
        invited_at=now_str,
        created_at=now_str
    )
    db.add(asg)

    sub = models.ReviewSubmission(
        id="sub-rep-1",
        assignment_id=asg.id,
        round_id=rnd.id,
        case_id=case.id,
        status="SUBMITTED",
        recommendation="MINOR_REVISION",
        total_weighted_score=8.75,
        created_at=now_str,
        updated_at=now_str
    )
    db.add(sub)

    c_public = models.ReviewComment(
        id="comm-rep-1",
        submission_id=sub.id,
        case_id=case.id,
        round_id=rnd.id,
        comment_type="AUTHOR_VISIBLE",
        comment_text="بحث متميز ومنهجية قوية. يرجى تدقيق المعادلة 3.",
        created_at=now_str
    )
    c_secret = models.ReviewComment(
        id="comm-rep-2",
        submission_id=sub.id,
        case_id=case.id,
        round_id=rnd.id,
        comment_type="CONFIDENTIAL_TO_EDITOR",
        comment_text="CONFIDENTIAL_EDITOR_SECRET_NOTE: الباحث متمكن والعمل جاهز للنشر بعد تعديلات بسيطة.",
        created_at=now_str
    )
    db.add_all([c_public, c_secret])

    # 6. Academic Profile
    prof = models.UnifiedAcademicProfile(
        id="prof-rep-601",
        organization_id=org_a_id,
        user_id=user_a_id,
        preferred_name_ar="د. سالم أحمد الأحمد",
        preferred_name_en="Dr. Salim A. Al-Ahmad",
        academic_title="دكتور",
        current_rank="PROFESSOR",
        department="قسم علوم الحاسب",
        completeness_score=95,
        created_at=now_str
    )
    db.add(prof)

    asset = models.ScholarlyAsset(
        id="asset-rep-1",
        organization_id=org_a_id,
        owner_user_id=user_a_id,
        title_ar="خوارزميات النمذجة المتقدمة للبيانات الأكاديمية",
        title_en="Advanced Modeling Algorithms for Academic Data",
        asset_type="JOURNAL_ARTICLE",
        publication_date="2025-01-01",
        journal_name="Journal of Academic Computing",
        doi="10.1016/j.jac.2025.01",
        created_at=now_str
    )
    db.add(asset)

    db.commit()


def test_research_project_report_export_json_docx_pdf():
    db = SessionLocal()
    user_a, org_a = create_test_tenant(db, "rep_salim", "org-rep-alpha", "RESEARCHER")
    create_test_tenant(db, "rep_editor", "org-rep-alpha", "ORGANIZATION_ADMIN")
    seed_report_test_entities(db, user_a.id, org_a.id, "usr-rep_editor")

    headers_a = get_auth_headers("rep_salim", org_a.id)

    # 1. JSON Export
    res_json = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "JSON",
        "language": "ar",
        "audience": "RESEARCHER"
    }, headers=headers_a)

    assert res_json.status_code == 200
    assert res_json.headers["Content-Type"].startswith("application/json")
    assert "X-Report-Integrity-Hash" in res_json.headers
    assert "X-Verification-Code" in res_json.headers

    data = res_json.json()
    assert data["schemaVersion"] == "1.0.0"
    assert data["reportType"] == "RESEARCH_PROJECT"
    assert data["title"]["ar"] == "أثر تقنيات الذكاء الاصطناعي في جودة التعليم العالي"
    assert len(data["sections"]) >= 4

    # 2. DOCX Export
    res_docx = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "DOCX",
        "language": "ar"
    }, headers=headers_a)

    assert res_docx.status_code == 200
    assert "wordprocessingml.document" in res_docx.headers["Content-Type"]
    assert len(res_docx.content) > 1000
    # Validate it's a valid ZIP/DOCX
    zip_f = zipfile.ZipFile(io.BytesIO(res_docx.content))
    assert "word/document.xml" in zip_f.namelist()

    # 3. PDF Export
    res_pdf = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "PDF",
        "language": "ar"
    }, headers=headers_a)

    assert res_pdf.status_code == 200
    assert res_pdf.headers["Content-Type"] == "application/pdf"
    assert res_pdf.content.startswith(b"%PDF")
    assert len(res_pdf.content) > 1000
    db.close()


def test_literature_synthesis_report_export():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")
    res = client.post("/api/reports/export", json={
        "report_type": "LITERATURE_SYNTHESIS",
        "source_id": "proj-rep-101",
        "format": "JSON",
        "language": "ar"
    }, headers=headers_a)

    assert res.status_code == 200
    data = res.json()
    assert data["reportType"] == "LITERATURE_SYNTHESIS"
    sec1 = data["sections"][0]
    assert sec1["keyMetrics"]["studies_count"] == 2


def test_prisma_flow_report_export():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")
    res = client.post("/api/reports/export", json={
        "report_type": "PRISMA_FLOW",
        "source_id": "proj-rep-101",
        "format": "JSON",
        "language": "ar"
    }, headers=headers_a)

    assert res.status_code == 200
    data = res.json()
    assert data["reportType"] == "PRISMA_FLOW"
    assert data["sections"][0]["tables"][0]["rows"][0][2] == 450


def test_promotion_readiness_report_export_and_disclaimer():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")
    res = client.post("/api/reports/export", json={
        "report_type": "PROMOTION_READINESS",
        "source_id": "app-rep-401",
        "format": "JSON",
        "language": "ar"
    }, headers=headers_a)

    assert res.status_code == 200
    data = res.json()
    assert data["reportType"] == "PROMOTION_READINESS"
    assert "أداة دعم قرار" in data["disclaimer"]["ar"]
    assert data["sections"][0]["keyMetrics"]["readiness_percentage"] == 98
    assert data["sections"][0]["keyMetrics"]["policy_version"] == 1


def test_peer_review_author_redaction_and_double_blind():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # Author requests review report -> strictly redacted
    res_author = client.post("/api/reports/export", json={
        "report_type": "PEER_REVIEW",
        "source_id": "case-rep-501",
        "format": "JSON",
        "language": "ar",
        "audience": "AUTHOR"
    }, headers=headers_a)

    assert res_author.status_code == 200
    data = res_author.json()
    serialized = json.dumps(data, ensure_ascii=False)

    # In Author report: Confidential comments & reviewer secret names MUST NOT BE PRESENT
    assert "CONFIDENTIAL_EDITOR_SECRET_NOTE" not in serialized
    assert "SECRET_REVIEWER_ALICE" not in serialized
    assert "CONFIDENTIAL_TO_EDITOR" not in serialized
    # Author visible comment SHOULD be present
    assert "بحث متميز ومنهجية قوية" in serialized


def test_peer_review_committee_full_export():
    headers_ed = get_auth_headers("rep_editor", "org-rep-alpha")

    # Committee / Editor requests report -> contains full review
    res_ed = client.post("/api/reports/export", json={
        "report_type": "PEER_REVIEW",
        "source_id": "case-rep-501",
        "format": "JSON",
        "language": "ar",
        "audience": "COMMITTEE"
    }, headers=headers_ed)

    assert res_ed.status_code == 200
    data = res_ed.json()
    serialized = json.dumps(data, ensure_ascii=False)
    assert "CONFIDENTIAL_EDITOR_SECRET_NOTE" in serialized
    assert "SECRET_REVIEWER_ALICE" in serialized


def test_peer_review_redaction_parity_across_json_docx_pdf():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # 1. Author JSON Export
    res_json = client.post("/api/reports/export", json={
        "report_type": "PEER_REVIEW",
        "source_id": "case-rep-501",
        "format": "JSON",
        "language": "ar",
        "audience": "AUTHOR"
    }, headers=headers_a)
    assert res_json.status_code == 200
    json_text = res_json.text
    assert "SECRET_REVIEWER_ALICE" not in json_text
    assert "CONFIDENTIAL_EDITOR_SECRET_NOTE" not in json_text

    # 2. Author DOCX Export
    res_docx = client.post("/api/reports/export", json={
        "report_type": "PEER_REVIEW",
        "source_id": "case-rep-501",
        "format": "DOCX",
        "language": "ar",
        "audience": "AUTHOR"
    }, headers=headers_a)
    assert res_docx.status_code == 200

    # Inspect all XML files inside the DOCX zip container
    zip_docx = zipfile.ZipFile(io.BytesIO(res_docx.content))
    for filename in zip_docx.namelist():
        xml_content = zip_docx.read(filename).decode("utf-8", errors="ignore")
        assert "SECRET_REVIEWER_ALICE" not in xml_content, f"Leakage found in DOCX {filename}"
        assert "CONFIDENTIAL_EDITOR_SECRET_NOTE" not in xml_content, f"Leakage found in DOCX {filename}"

    # 3. Author PDF Export
    res_pdf = client.post("/api/reports/export", json={
        "report_type": "PEER_REVIEW",
        "source_id": "case-rep-501",
        "format": "PDF",
        "language": "ar",
        "audience": "AUTHOR"
    }, headers=headers_a)
    assert res_pdf.status_code == 200
    pdf_bytes = res_pdf.content
    assert b"SECRET_REVIEWER_ALICE" not in pdf_bytes
    assert b"CONFIDENTIAL_EDITOR_SECRET_NOTE" not in pdf_bytes


def test_report_audience_privilege_escalation_blocked():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # Author attempts privilege escalation by demanding COMMITTEE view
    res_escalate_review = client.post("/api/reports/export", json={
        "report_type": "PEER_REVIEW",
        "source_id": "case-rep-501",
        "format": "JSON",
        "language": "ar",
        "audience": "COMMITTEE"
    }, headers=headers_a)
    assert res_escalate_review.status_code == 403

    # Applicant attempts privilege escalation for promotion report
    res_escalate_prom = client.post("/api/reports/export", json={
        "report_type": "PROMOTION_READINESS",
        "source_id": "app-rep-401",
        "format": "JSON",
        "language": "ar",
        "audience": "COMMITTEE"
    }, headers=headers_a)
    assert res_escalate_prom.status_code == 403


def test_academic_profile_report_export():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")
    res = client.post("/api/reports/export", json={
        "report_type": "ACADEMIC_PROFILE",
        "source_id": "prof-rep-601",
        "format": "JSON",
        "language": "ar"
    }, headers=headers_a)

    assert res.status_code == 200
    data = res.json()
    assert data["reportType"] == "ACADEMIC_PROFILE"
    assert data["sections"][0]["keyMetrics"]["publications_count"] == 1


def test_cross_tenant_report_export_isolation():
    db = SessionLocal()
    create_test_tenant(db, "rep_mona", "org-rep-beta", "RESEARCHER")
    db.close()
    headers_b = get_auth_headers("rep_mona", "org-rep-beta")

    # User in Org B tries to export Project in Org A -> 404
    res = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "JSON"
    }, headers=headers_b)

    assert res.status_code == 404


def test_same_tenant_unauthorized_user_export():
    db = SessionLocal()
    create_test_tenant(db, "rep_unauthorized_peer", "org-rep-alpha", "MEMBER")
    db.close()
    headers_peer = get_auth_headers("rep_unauthorized_peer", "org-rep-alpha")

    # Non-admin, non-owner tries to export someone else's project -> 403 Forbidden
    res = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "JSON"
    }, headers=headers_peer)

    assert res.status_code == 403


def test_canonical_report_context_endpoint():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    res = client.post("/api/reports/context", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "language": "ar",
        "audience": "RESEARCHER"
    }, headers=headers_a)

    assert res.status_code == 200
    ctx = res.json()
    assert ctx["manifest"]["report_type"] == "RESEARCH_PROJECT"
    assert "sections" in ctx
    assert len(ctx["sections"]) >= 4

    # Privilege escalation on context endpoint must also fail with 403
    res_bad = client.post("/api/reports/context", json={
        "report_type": "PEER_REVIEW",
        "source_id": "case-rep-501",
        "language": "ar",
        "audience": "COMMITTEE"
    }, headers=headers_a)
    assert res_bad.status_code == 403


def test_cross_format_data_parity_and_hashes():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # Generate JSON, DOCX, and PDF from the same canonical source
    res_json = client.post("/api/reports/export", json={
        "report_type": "PROMOTION_READINESS",
        "source_id": "app-rep-401",
        "format": "JSON",
        "language": "ar"
    }, headers=headers_a)

    res_docx = client.post("/api/reports/export", json={
        "report_type": "PROMOTION_READINESS",
        "source_id": "app-rep-401",
        "format": "DOCX",
        "language": "ar"
    }, headers=headers_a)

    res_pdf = client.post("/api/reports/export", json={
        "report_type": "PROMOTION_READINESS",
        "source_id": "app-rep-401",
        "format": "PDF",
        "language": "ar"
    }, headers=headers_a)

    assert res_json.status_code == 200
    assert res_docx.status_code == 200
    assert res_pdf.status_code == 200

    # Validate that the document hash returned in headers exactly matches the sha256 of delivered bytes
    hash_json = res_json.headers["X-Report-Integrity-Hash"]
    hash_docx = res_docx.headers["X-Report-Integrity-Hash"]
    hash_pdf = res_pdf.headers["X-Report-Integrity-Hash"]

    assert hashlib.sha256(res_json.content).hexdigest() == hash_json
    assert hashlib.sha256(res_docx.content).hexdigest() == hash_docx
    assert hashlib.sha256(res_pdf.content).hexdigest() == hash_pdf


def test_bilingual_and_arabic_rtl_rendering():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # Generate Bilingual and English reports
    res_bi = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "PDF",
        "language": "bilingual"
    }, headers=headers_a)

    res_en = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "DOCX",
        "language": "en"
    }, headers=headers_a)

    assert res_bi.status_code == 200
    assert res_bi.headers["Content-Type"] == "application/pdf"
    assert res_en.status_code == 200
    assert "wordprocessingml.document" in res_en.headers["Content-Type"]


def test_report_verification_endpoint():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # Generate a report to register verification code
    res = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "JSON"
    }, headers=headers_a)

    verif_code = res.headers["X-Verification-Code"]

    # Verify valid code
    res_verif = client.get(f"/api/reports/verify/{verif_code}")
    assert res_verif.status_code == 200
    data = res_verif.json()
    assert data["valid"] is True
    assert data["verification_code"] == verif_code
    assert data["report_type"] == "RESEARCH_PROJECT"

    # Public verification should never leak internal user IDs, secret tokens, or full reports
    serialized = json.dumps(data)
    assert "user_id" not in serialized
    assert "download_token" not in serialized
    assert "password" not in serialized

    # Verify bogus code
    res_bogus = client.get("/api/reports/verify/BSR-FAKE-0000")
    assert res_bogus.status_code == 200
    assert res_bogus.json()["valid"] is False

    # Verify malformed code
    res_malformed = client.get("/api/reports/verify/INVALID$$$CODE@@")
    assert res_malformed.status_code == 200
    assert res_malformed.json()["valid"] is False


def test_invalid_report_types_formats_and_languages():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # Invalid report type
    res1 = client.post("/api/reports/export", json={
        "report_type": "INVALID_TYPE",
        "source_id": "proj-rep-101",
        "format": "JSON"
    }, headers=headers_a)
    assert res1.status_code == 400

    # Invalid format
    res2 = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "EXE"
    }, headers=headers_a)
    assert res2.status_code == 400

    # Invalid language
    res3 = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "JSON",
        "language": "fr"
    }, headers=headers_a)
    assert res3.status_code == 400


def test_docx_and_pdf_structural_validity():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")

    # Check DOCX validity with python-docx
    res_docx = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "DOCX",
        "language": "ar"
    }, headers=headers_a)
    assert res_docx.status_code == 200

    doc = docx.Document(io.BytesIO(res_docx.content))
    assert len(doc.paragraphs) > 5
    assert len(doc.tables) >= 2

    # Check PDF validity
    res_pdf = client.post("/api/reports/export", json={
        "report_type": "RESEARCH_PROJECT",
        "source_id": "proj-rep-101",
        "format": "PDF",
        "language": "ar"
    }, headers=headers_a)
    assert res_pdf.status_code == 200
    assert res_pdf.content.startswith(b"%PDF-")
    assert b"%%EOF" in res_pdf.content[-2048:]


def test_mutation_free_generation():
    headers_a = get_auth_headers("rep_salim", "org-rep-alpha")
    db = SessionLocal()

    # Capture state before export
    proj_before = db.query(models.ResearchProject).filter_by(id="proj-rep-101").first()
    title_before = proj_before.titleAr

    # Generate multiple reports
    client.post("/api/reports/export", json={"report_type": "RESEARCH_PROJECT", "source_id": "proj-rep-101", "format": "JSON"}, headers=headers_a)
    client.post("/api/reports/export", json={"report_type": "RESEARCH_PROJECT", "source_id": "proj-rep-101", "format": "PDF"}, headers=headers_a)

    # Verify no domain entity fields mutated
    proj_after = db.query(models.ResearchProject).filter_by(id="proj-rep-101").first()
    assert proj_after.titleAr == title_before
    db.close()
